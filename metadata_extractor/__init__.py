from collections import Counter, defaultdict
from converter import save_metadata
from docx import Document
from docx.shared import RGBColor
from numpy import median, std
from os.path import dirname, isfile, realpath
from re import finditer, MULTILINE, search
from requests import get
from string import Template
from xml.etree import ElementTree

PATH_TO_DIRECTORY_OF_THIS_FILE = dirname(realpath(__file__))

with open(PATH_TO_DIRECTORY_OF_THIS_FILE + "/prep/stopwords.txt") as f:
    stopwords = f.read().strip().split("\n")
with open(PATH_TO_DIRECTORY_OF_THIS_FILE + "/prep/tags.txt") as f:
    tags = sorted(f.read().strip().split("\n"), key=lambda tag: -1*len(tag))

def flatten(iterable, levels=100):
    flattened = []
    levels -= 1
    for item in iterable:
        if isinstance(item, list) and levels >= 0:
            flattened.extend(flatten(item, levels))
        else:
            flattened.append(item)
    return flattened
    
def simplify_positions(positions):
    #positions = [3, 4, 5, 6, 10, 11, 13, 14, 15, 16, 20, 21, 22, 23, 24, 25, 28, 29, 30, 31]
    #positions = [3, 6, 10, 11, 13, 16, 20, 25, 28, 31]
    return [position for position in positions if position - 1 not in positions or position + 1 not in positions]
    

def join_nested_items(_input, joiner="\n"):
    text = ""
    if isinstance(_input, unicode) or isinstance(_input, str):
        text = _input
    elif isinstance(_input, list):
        for item in _input:
            if isinstance(item, list):
                text += join_nested_items(item)
            elif isinstance(item, unicode) or isinstance(item, str):
                text += joiner + item
    return text

def count_nested_items(iterable):
    count = 0 
    for item in iterable:
        if isinstance(item, list):
            count += count_nested_items(item)
        else:
            count += 1
    return count

# not splitting this correctly!!!! need to figure out offsetting even with iterables
# or maybe not returning stuff with correct depth
def split_by_indices(items, positions, offset=0, return_i=False):
    #last_position = len(flatten(items))
    offset_positions = [position - offset for position in positions]
    #if 2 in positions and 27 in positions and 54 in positions: 
    #print "\n\nstarting split_by_indices with " 
    #print "\titerable:", items
    #print "\tpositions:", positions
    #print "\toffset:", offset
    #print "\toffset_positions:", offset_positions
    try:
        #result = [item for item in [iterable[offset_positions[current_index - 1] if current_index > 0 else 0 : current_position] for current_index, current_position in enumerate(offset_positions)] if item]
        result = []
        sublist = []
        i = 0
        for item in items:
            if isinstance(item, list):
                if sublist:
                    result.append(sublist)
                item, i = split_by_indices(item, positions, offset=i, return_i=True)
                result.append(item)
            else:
                if i in offset_positions:
                    result.append(sublist)
                    sublist = [item]
                else:
                    sublist.append(item)
                i += 1
        if sublist: result.append(sublist)
        if return_i:
            return result, i
        else:
            return result
        
    except Exception as e:
        print "EXCEPTION trying to split by indices:", e

# should I lowercase everything??                                                                                                                               
def tokenize(text):                                                                                                                                             
    punctuation = ["%3A","%3B","%3C","%3D","%3E","%3F","{","}","-","_","/",":","_",".","\\",">","<","&amp;","=","+","'",'"',"&","#","%2C",")","(","@","?",":","!",";",","]                                                                                                                                                      
    for p in punctuation:                                                                                                                                       
        text = text.replace(p, " ")                                                                                                                             
    terms = text.split()
    terms = [term for term in terms if 3 < len(term) < 10 and not search("\d", term) and term not in stopwords]
    return terms                                          

def is_metadata(inpt, debug=False):
    if debug: print "starting is_metadata with", type(inpt)
    _type = str(type(inpt))
    if isinstance(inpt, str) or isinstance(inpt, unicode):
        if isfile(inpt):
            if inpt.endswith(".txt"):
                with open(inpt) as f:
                    text = f.read()
                return extract_metadata_from_text(text)
            elif inpt.endswith(".doc") or inpt.endswith(".docx"):
                with open(inpt) as f:
                    text = "\n".join([p.text for p in Document(f).paragraphs])
        elif inpt.strip().startswith("http") and inpt.strip().count("\n") < 3:
            if debug: print "inpt is a url"
            text = get(inpt).text
        else:
            if debug: print "inpt is regular text"
            text = inpt
    elif _type == "<class 'django.core.files.uploadedfile.InMemoryUploadedFile'>":
        text = "\n".join([p.text.encode("utf-8") for p in Document(inpt).paragraphs])

    tag_parts = flatten([tag.split() for tag in tags])
    tokens = tokenize(text)
    percentage = float(len([token for token in tokens if token in tag_parts])) / len(tokens)
    if debug: print "percentage:", percentage

    # if more than 5% of the input is metadata terms, it's about metadata
    return percentage > 0.05
                
def extract_metadata(inpt):
    _type = str(type(inpt))
    if isinstance(inpt, str) or isinstance(inpt, unicode):
        if isfile(inpt):
            print "filepath"
            if inpt.endswith(".txt"):
                with open(inpt) as f:
                    text = f.read()
                return extract_metadata_from_text(text)
            elif inpt.endswith(".doc") or inpt.endswith(".docx"):
                with open(inpt) as f:
                    doc = Document(f)
                return extract_metadata_from_doc(doc)
    elif _type == "<class 'django.core.files.uploadedfile.InMemoryUploadedFile'>":
        doc = Document(inpt)
        return extract_metadata_from_doc(doc)




def extract_metadata_from_doc(doc, debug=False):

    tree = treeify(doc)

    #print "tree:", type(tree)

    level = 0
    while level <= 10:
        #print "level:", level
        counters = []
        items = flatten(tree, levels=level)
        for item in items:
            text = join_nested_items(item)
            #print "\tlen text:", len(text)
            counter = Counter()
            for tag in tags:
                counter[tag] += text.count(tag)
                #count = counter[tag]
                #if count > 0:
                #    print "\t",tag,":", count
            #print "\t\tcounter:", counter.most_common(5)
            counters.append(counter)

        averages = Counter()
        more_than_two = 0
        for tag in tags:
            averages[tag] = average = median([counter[tag] for counter in counters])
            if average > 2:
                more_than_two += 1

        print "more_than_two:", more_than_two
        print "averages:", averages.most_common(5)
        if more_than_two < 3:
            print "breaking on level", level
            break
        else:
            level += 1

    if debug:
        for level in [0,1,2,3,4]:
            with open("/tmp/level_" + str(level) + ".txt", "wb") as f:
                for i, item in enumerate(flatten(tree, levels=level)):
                   text = join_nested_items(item)
                   f.write("\n" + "|*|" + str(i) + "|*|" + text.encode("utf-8"))

    metadata = [] 
    for item in items:
        for d in parse_text(join_nested_items(item)):
            if d:
                metadata.append(d)
        
    #print "metadata:", metadata 
    return metadata
  
# only works if one layer in text now 
def parse_text(text):

    print "starting parse_text with:", type(text), len(text)
   
    layers = []
    layer = {}

    hit_tag = False
    for line in text.split("\n"):
        for tag in tags:
            if tag in line:
                hit_tag = True
                break
        if hit_tag:
            break
        if line and 5 < len(line) < 50:
            name = line.strip()
            #print "name:", name
            layer['NAME'] = name

    pattern = "^(?P<key>" + "|".join([tag.replace("(","\(").replace(")","\)") for tag in tags]) + ")[(\r\n)|\n](?P<value>[^\n]+)$"
    #print "pattern:", pattern
    #print "[]" * 30
    for mg in finditer(pattern, text, MULTILINE):
        try: print mg.group("key"), ":>>:", mg.group("value")
        except: pass
        layer[mg.group("key")] = mg.group("value")
    #print "\nlayer:", layer
    layers.append(layer)
    """
    active_tag = None
    active_text = ""
    for p in text.split("\n"):
        for tag in tags:
            if tag in p:
                if tag in layer:
                    layers.append(layer)
                    layer = {}
                    active_tag = tag
                else:
                    if active_text:
                        layer[tag] = active_text
                        print "tag:", tag
                    active_tag = tag
            else:
                active_text += p
    """

    return layers

def treeify(doc):

    d = {}

    paragraphs = doc.paragraphs
    number_of_paragraphs = len(paragraphs)

    for i, p in enumerate(paragraphs):
        bold = p.style.font.bold or False
        color = p.style.font.color.rgb or RGBColor(0x00, 0x00, 0x00)
        italic = p.style.font.italic or False
        size = (p.style.font.size and p.style.font.size.pt) or (doc.styles['Normal'].font.size and doc.styles['Normal'].font.size.pt) or 12
        key = (bold, color, italic, size)
        if len(p.text) > 5:
            if key in d:
                d[key]['positions'].append(i)
            else:
                d[key] = {"bold": bold, "positions": [i], "size": size}

    for key in d:
        positions = d[key]['positions']
        d[key]['positions_for_splitting'] = simplify_positions(positions)
        d[key]['count'] = len(positions)
        d[key]['std'] = std(positions) / number_of_paragraphs

    # filter out styles that don't have an even distribution, which is defined as std of at least 25%
    d = dict([(key, value) for key, value in d.items() if value['std'] > 0.25])

    for k in d:
        v = d[k]
        print "k:", k
        print "\tcount:", v['count']
        print "\tsize:", v['size']
        print "\tstd:", v['std']

    bolds = []
    counts = []
    sizes = []
    stds = []
    for key, value in d.items():
        bolds.append(value['bold'])
        counts.append(value['count'])
        sizes.append(value['size'])
        stds.append(value['std'])

    bolds.sort()
    bolds.reverse()

    counts.sort()

    sizes.sort()
    sizes.reverse()

    stds.sort()
    stds.reverse()

    print "\n"
    print "bolds:", bolds
    print "counts:", counts
    print "sizes:", sizes
    print "stds:", stds
  
    for key, value in d.items():
        # multiplying std by 0.9 makes it a little less important in tie-breaking situations
        value['score'] = bolds.index(value['bold']) + counts.index(value['count']) + stds.index(value['std']) * 0.9 + sizes.index(value['size'])

    styles = sorted(d.items(), key=lambda t: t[1]['score'])

    for level, (key, style) in enumerate(styles):
        print "\n\nlevel:", level
        print "key:", key
        print "bold:", style['bold']
        print "count:", style['count']
        print "positions:", style['positions'][:20]
        print "positions_for_splitting:", style['positions_for_splitting'][:20], "..."
        print "size:", style['size']
        print "std:", style['std']
        print "score:", style['score']

    paragraphs = [p.text for p in paragraphs]
    #tree = split_by_indices(paragraphs, styles[0][1]['positions_for_splitting'])
    # xxx
    tree = split_by_indices(paragraphs, [p-1 for p in styles[0][1]['positions_for_splitting']])
    for key, style in styles[1:]:
        print "\nstyle:", key
        #tree = [split_by_indices(item, style['positions_for_splitting'], offset=sum(map(count_nested_items, tree[0:i]))) for i, item in enumerate(tree)]
        # xxx
        tree = [split_by_indices(item, [p-1 for p in style['positions_for_splitting']], offset=sum(map(count_nested_items, tree[0:i]))) for i, item in enumerate(tree)]
        print "tree:", count_nested_items(tree)

    return tree


def extract_metadata_from_text(text):
    print "starting extract_metadata_from_text with ", len(text)

    pieces = get_pieces(text)

def get_pieces(text):
    print "starting get_pieces with", len(text)
    #paragraphs = [p for p in text.split("\n") if len(p) >= 3]
    paragraphs = [p for p in text.split("\n")]
    print "paragraphs = ", paragraphs[:10]
    number_of_docs = len(paragraphs)
    print "number_of_docs:", number_of_docs

    # do tfidf
    document_counts = Counter()
    term_counts = []
    #term_frequencies = []
    pieces = []
    for paragraph in paragraphs:
        piece = {}
        piece['text'] = paragraph
        piece['tokens'] = tokens = tokenize(paragraph)
        document_counts.update(set(tokens))
        piece['number_of_tokens'] = number_of_tokens = len(tokens)
        piece['term_count'] = term_count = Counter(tokens)
        term_counts.append(term_count)
        piece['term_frequency'] = dict([ ( token, float(count) / number_of_tokens ) for token, count in term_count.items() ])
        pieces.append(piece)

    document_frequency = dict([(term, float(count) / number_of_docs) for term, count in document_counts.items()])
    #print "document_frequency:", document_frequency

    for piece in pieces:
        piece['term_tfidf'] = term_tfidf = sorted([(term, float(frequency) / document_frequency[term]) for term, frequency in piece['term_frequency'].items()], key=lambda t: -1*t[1])
        piece['terms'] = [term for term, tfidf in term_tfidf]

    pieces = pieces

    def find_and_merge_pieces(pieces):
        #print "starting find_and_merge_pieces with", len(pieces), "pieces"
        print ".",
        #raw_input("press enter to continue")
        number_of_pieces = len(pieces)
        for i in range(number_of_pieces):
            current_piece = pieces[i]
            #print "current_piece:", current_piece
            #raw_input("press enter to continue")
            if i == 0:
                previous_piece = current_piece
            else:
                #print "\n\n"
                #print "text:", current_piece['text']
                #print "terms:", current_piece['terms']
                combined_term_count = previous_piece['term_count'] + current_piece['term_count']
                #print "combined_term_count:", combined_term_count
                number_of_terms = sum(combined_term_count.values())
                combined_term_frequency = dict([(term, float(count) / number_of_terms) for term, count in combined_term_count.items()])
                combined_term_tfidf = sorted([(term, float(frequency) / document_frequency[term]) for term, frequency in combined_term_frequency.items()], key=lambda t: -1*t[1])
                combined_terms = [term for term, tfidf in combined_term_tfidf]
                number_of_combined_terms = len(combined_terms)
                #print "previous_terms = ", previous_piece['terms']
                #print "combined terms are", combined_terms
                #print "x:", [term in combined_terms[:8] for term in previous_piece['terms'][:4]]
                number_of_previous_terms = len(previous_piece['terms'])
                if number_of_previous_terms == 0 or float(sum([term in combined_terms[:4] for term in previous_piece['terms'][:3]])) / len(previous_piece['terms']) >= 0.75:
                    #print "merging"
                    previous_piece['text'] += "\n" + current_piece['text']
                    previous_piece['term_count'] = combined_term_count
                    previous_piece['term_frequency'] = combined_term_frequency
                    previous_piece['terms'] = combined_terms
                    previous_piece['term_tfidf'] = combined_term_tfidf
                    pieces.remove(current_piece) # only specific one or all similar
                    return True

                previous_piece = current_piece
    while find_and_merge_pieces(pieces):
        pass

    pieces.reverse()
    print "len(pieces):", len(pieces)

    while find_and_merge_pieces(pieces):
        pass


    pieces.reverse()
    print "len(pieces):", len(pieces)

    with open("/tmp/output.txt", "wb") as f:
        f.write(("\n" + "=" * 25 + "\n").join([p['text'].encode("utf-8") for p in pieces]))
